
import os

import re
import jieba

import collections
import codecs
from util.operate_mysql import *
from util.operate_mysql import clear_table as clear_table
import docx
import doc
import xlrd
import xlwt
from xlutils.copy import copy

docx_path_56 = 'F:\\pythoncode\\score\\dst_docx_56\\' #要分析的文件目录
student_idname_56 = 'F:\\pythoncode\\score\\jk152056.xls'
score_fiel_56 = 'F:\\pythoncode\\score\\jk152056score.xls'

docx_path_78 = 'F:\\pythoncode\\score\\dst_docx_78\\' #要分析的文件目录
student_idname_78 = 'F:\\pythoncode\\score\\jk152078.xls'
score_fiel_78 = 'F:\\pythoncode\\score\\jk152078score.xls'

def get_max_min_size(docx_path):
    min = 0
    max = 0
    file_list = os.listdir(docx_path)
    for file_name in file_list:
        file_name_list = file_name.split(".")
        if file_name_list[1] == "docx":
            document = docx.Document(docx_path+file_name)
        #elif file_name_list[1] == "doc":
        #    document = doc.Document(docx_path+file_name)
        contents_list = []
        # 取文件内容
        for paragraph in document.paragraphs:
            contents_list.append(paragraph.text.encode('utf-8'))
        # 分词
        words_list = []
        file_len = 0
        for text in contents_list:
            file_len = file_len + len(text)

        if file_len > max:
            max = file_len
        if 0 == min:
            min = file_len
        if file_len<min:
            min = file_len
    return max, min


def get_score_from_file(file_name, max, min):
    document = docx.Document(file_name)
    contents_list = []
    # 取文件内容
    for paragraph in document.paragraphs:
        contents_list.append(paragraph.text.encode('utf-8'))
    # 按文件长度计分
    words_list = []
    file_len = 0
    for text in contents_list:
        file_len = file_len + len(text)

    score = round(file_len/max * 15) + 80
    return score



def parase_all_and_record_score(docx_path,student_idname,score_fiel):
    #打开学生名单excel，copy一份另存并写入成绩
    source_excel_file = xlrd.open_workbook(student_idname)
    dest_excel_file = copy(source_excel_file)
    read_table = source_excel_file.sheet_by_name(u'Sheet1')  # 通过名称获取
    write_table = dest_excel_file.get_sheet(0)

    max, min = get_max_min_size(docx_path)
    file_list = os.listdir(docx_path)

    #初始化所有人成绩为0
    for i in range(read_table.nrows):
        get_name_fromcell = read_table.cell(i, 2).value
        write_table.write(i, 3, 0)

    for file_name in file_list:
        file_name_list = []
        file_name_list = file_name.split('-')
        student_tmp = file_name_list[2].split(".")
        student_name = student_tmp[0]
        score = 0
        for i in range(read_table.nrows):
            get_name_fromcell = read_table.cell(i,2).value
            if(get_name_fromcell == student_name):
                print(student_tmp)
                if student_tmp[1] == "docx":
                    score = get_score_from_file(docx_path+file_name, max, min)
                #elif student_tmp == "doc":
                #    document = doc.Document(file_name)
                #    score = get_score_from_file(document, max, min)
                write_table.write(i, 3, score)

    dest_excel_file.save(score_fiel)

def parase_all_and_insert_db(docx_path):
    clear_table('words_temp')  # 清除数据
    operatMySQl = OperateMySQL()

    jieba.load_userdict("dict.txt")

    # 获取文本内容
    contents_list = []
    words_list = []
    file_list = os.listdir(docx_path)
    for filename in file_list:
        filepathname= docx_path + filename
        print(filepathname)
        contents_list.clear()
        document = docx.Document(filepathname)
        #取文件内容
        for paragraph in document.paragraphs:
            contents_list.append(paragraph.text.encode('utf-8'))
        # 分词
        for text in contents_list:
            splits = jieba.cut(text)
            for word in splits:
                words_list.append(word.strip().lower())

    # 统计分词频率，计入数据库
    counter = collections.Counter(words_list)

    # 插入数据库
    for collect in counter:
        sqli = "insert into words_temp values ('{0}',{1});"
        sqlm = sqli.format(collect, counter[collect])
        operatMySQl.execute(sqlm)
    operatMySQl.commit()


    # 清洗无用数据
    clean_date_table = [ ' ', '\t', '\n', '、', '.', '，', '/', '；', '（', '）','。','：',',','-',';',')',':',
                        '的', '和','熟悉','开发','等','有','相关','及','能力','优先','工作','良好','与','对','年','常用',
                        '能够','了解','者','专业','熟练','能','计算机','专业','电子','进行',
                        '以上','具有','使用','精通','或','嵌入式','工程师',
                        ' ','1','2','3','4','5','6','7','8','9','0']
    sqli = "delete from words_temp  where words_content = '{0}';"
    for clean_item in clean_date_table:
        sqlm = sqli.format(clean_item)
        operatMySQl.execute(sqlm)
    operatMySQl.commit()


if __name__ == '__main__':
    starttime = datetime.datetime.now()
    print('Start time is %s.' % (str(datetime.datetime.now())))

    parase_all_and_insert_db(docx_path_56)

    parase_all_and_record_score(docx_path_56,student_idname_56,score_fiel_56)

    #parase_all_and_insert_db(docx_path_78)

    #parase_all_and_record_score(docx_path_78,student_idname_78,score_fiel_78)

    # 程序结束时间 及 耗时
    timedelta = datetime.datetime.now() - starttime
    print('End time is %s.' % (str(datetime.datetime.now())))
    print('Total test execution duration is %s.' % (timedelta.__str__()))



