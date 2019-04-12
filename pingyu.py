
import os

import re
import jieba

import collections
import codecs
from util.operate_mysql import *
from util.operate_mysql import clear_table as clear_table
import docx
#import doc
import xlrd
import xlwt
from xlutils.copy import copy
from docx.oxml.ns import qn

docx_path_56 = 'F:\\pythoncode\\score\\dst_docx_56\\' #要分析的文件目录
student_idname_56 = 'F:\\pythoncode\\score\\jk152056.xls'
score_fiel_56 = 'F:\\pythoncode\\score\\jk152056score.xls'

docx_path_78 = 'F:\\pythoncode\\score\\dst_docx_78\\' #要分析的文件目录
student_idname_78 = 'F:\\pythoncode\\score\\jk152078.xls'
score_fiel_78 = 'F:\\pythoncode\\score\\jk152078score.xls'

docx_path_all = 'F:\\pythoncode\\score\\dst_docx_all\\'

pingyu_56  = 'F:\\pythoncode\\score\\pingyu_56\\' #评语输出目录
score_pingyu_56 = 'F:\\pythoncode\\score\\score_pingyu_56.xls'
score_pingyu_78 = 'F:\\pythoncode\\score\\score_pingyu_78.xls'

#score_pingyu_56 = 'F:\\pythoncode\\score\\yytmp.xls'
#pingyu_56  = 'F:\\pythoncode\\score\\yypingyu\\' #评语输出目录

def get_max_min_size(docx_path):
    min = 0
    max = 0
    file_list = os.listdir(docx_path)
    for file_name in file_list:
        file_name_list = file_name.split(".")
        if file_name_list[1] == "docx":
            document = docx.Document(docx_path+file_name)
        #elif file_name_list[1] == "doc":
        #    document = doc.Document(docx_path+file_name)
        contents_list = []
        # 取文件内容
        for paragraph in document.paragraphs:
            contents_list.append(paragraph.text.encode('utf-8'))
        # 分词
        words_list = []
        file_len = 0
        for text in contents_list:
            file_len = file_len + len(text)

        if file_len > max:
            max = file_len
        if 0 == min:
            min = file_len
        if file_len<min:
            min = file_len
    return max, min


def get_score_from_file(file_name, max, min):
    document = docx.Document(file_name)
    contents_list = []
    # 取文件内容
    for paragraph in document.paragraphs:
        contents_list.append(paragraph.text.encode('utf-8'))
    # 按文件长度计分
    words_list = []
    file_len = 0
    for text in contents_list:
        file_len = file_len + len(text)

    score = round(file_len/max * 15) + 80
    return score

def insert_pingyu(file_name,input_path, outpu_path, score):
    document = docx.Document(input_path+file_name)
    tables = document.tables
    try:
        if(score.isdigit() == True):
            score = int(score,10)
        else:
            print(file_name)
            return
    except:
        pass

    lasttable = None
    firsttable = None
    for table in tables:
        if firsttable is None:
            firsttable = table
        lasttable = table
    '''
    fm_cells = firsttable.rows[0].cells
    fm_cells[1].paragraphs[0].style.font.name = '黑体'
    fm_cells[1].paragraphs[0].style.font.size = 155000  # 小四号字体
    fm_cells[1].text = '项目'
    '''

    pingyu = ''

    if score >= 10 and score < 60:
        pingyu = '该生在实训过程中，态度较较差，对团队工作贡献较少，未达到小学期的实训要求。'
    elif score >= 60 and score < 65:
        pingyu = '该生在实训过程中，态度一般，对团队工作贡献较为一般，所搭建系统完成了项目要求的功能，基本达到了实训要求。'
    elif score >= 65 and score < 70:
        pingyu = '该生在实训过程中，态度一般，在团队中做了一定的工作，所搭建系统完成了项目要求的功能，基本达到了实训要求。'
    elif score >= 70 and score < 80:
        pingyu = '该生在实训过程中，态度认真，在团队中做了一定的工作，所搭建系统完成了大部分项目要求的功能，达到了实训要求。'
    elif score >= 80 and score < 90:
        pingyu = '该生在实训过程中，态度比较认真，对团队贡献较大，所搭建系统完成了所有项目要求的功能，并完成了部分增强功能，较好的达到了实训要求。'
    elif score >= 90 and score <= 100:
        pingyu = '该生在实训过程中，态度非常认真，积极帮助其他同学，所搭建系统完成了所有项目要求的功能，并完成了部分增强功能，很好的达到了实训要求。'
    else:
        print("没有学生成绩：", file_name, score)


    ''' #yueyue的评语
    if score >= 60 and score < 70:
        pingyu = '该生在小学期实训过程中，态度较认真，但对团队工作贡献一般，拓扑搭建完成了连通的基本功能，基本达到了小学期的实训要求。'
    elif score >= 70 and score < 75:
        pingyu = '该生在小学期实训过程中，态度认真，在团队中做了一定的工作，拓扑搭建完成了部分增强功能，达到了小学期的实训要求。'
    elif score >= 75 and score < 80:
        pingyu = '该生在小学期实训过程中，态度认真，对团队贡献较大，拓扑搭建完成了大部分功能，较好的达到了小学期的实训目的。'
    elif score >= 80:
        pingyu = '该生在小学期实训过程中，态度认真，积极帮助其他同学，拓扑搭建基本完成了所有功能，很好的达到了小学期的实训目的。'
    else:
        print("没有学生成绩：", file_name, score)
    '''
    #'''
    hdr_cells = lasttable.rows[0].cells
    #lasttable.cell(0, 0).paragraphs[0].style.font.name = 'liguofu'
    #lasttable.cell(0, 0).paragraphs[0].style.font.color.rgb = RGBColor(22, 120, 190)
    #hdr_cells[0].paragraphs[0].style.font.name = 'liguofu'

    #document.styles['Normal'].font.name = u'liguofu'
    #document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'liguofu')

    #hdr_cells[1].paragraphs[0].style.font.name = u'liguofu'
   # hdr_cells[1].paragraphs[0].style._element.rPr.rFonts.set(qn('w:eastAsia'), u'liguofu')

    hdr_cells[1].paragraphs[0].style.font.size = 155000 #小四号字体
    hdr_cells[1].text = pingyu

    hdr_cells = lasttable.rows[1].cells
    hdr_cells[1].text = str(int(score))

    hdr_cells = lasttable.rows[2].cells
    #hdr_cells[1].text = '李楸桐'
    hdr_cells[3].text = '2018年12月3日'
    #'''

    #print(len(lasttable.rows))
    #print(len(lasttable.columns))
    document.save(outpu_path+file_name)

#将excel中成绩写入对应word
def parase_all_and_insert_pingyu(docx_path,outpu_path,score_pingyu):
    # 打开学生成绩excel
    source_excel_file = xlrd.open_workbook(score_pingyu)
    read_table = source_excel_file.sheet_by_name(u'Sheet1')  # 通过名称获取

    file_list = os.listdir(docx_path)
    for file_name in file_list:
        #file_name_list = []
        file_name_list = file_name.split('.')
        student_tmp = file_name_list[0].split('-')
        student_name = student_tmp[1]
        score = 0
        find = 0
        banji = 0
        xuehao = ''
        for i in range(read_table.nrows):
            get_name_fromcell = read_table.cell(i, 2).value
            if (get_name_fromcell == student_name):
                # print(student_tmp)
                score = read_table.cell(i, 3).value # 取成绩
                xuehao = read_table.cell(i,1).value # 取成绩
                banji = read_table.cell(i,0).value # 取成绩
                find = 1
        if (find == 0):
            print("没有学生：", file_name, student_name)
        else:
            insert_pingyu(file_name, docx_path, outpu_path, score)
            '''
            if banji == 6:
                pingyu = ''

                if score >= 10 and score < 60:
                    pingyu = '该生在实训过程中，态度较较差，对团队工作贡献较少，未达到小学期的实训要求。'
                elif score >= 60 and score < 65:
                    pingyu = '该生在实训过程中，态度一般，对团队工作贡献较为一般，所搭建系统完成了项目要求的功能，基本达到了实训要求。'
                elif score >= 65 and score < 70:
                    pingyu = '该生在实训过程中，态度一般，在团队中做了一定的工作，所搭建系统完成了项目要求的功能，基本达到了实训要求。'
                elif score >= 70 and score < 80:
                    pingyu = '该生在实训过程中，态度认真，在团队中做了一定的工作，所搭建系统完成了大部分项目要求的功能，达到了实训要求。'
                elif score >= 80 and score < 90:
                    pingyu = '该生在实训过程中，态度比较认真，对团队贡献较大，所搭建系统完成了所有项目要求的功能，并完成了部分增强功能，较好的达到了实训要求。'
                elif score >= 90 and score <= 100:
                    pingyu = '该生在实训过程中，态度非常认真，积极帮助其他同学，所搭建系统完成了所有项目要求的功能，并完成了部分增强功能，很好的达到了实训要求。'
                else:
                    print("没有学生成绩：", file_name, score)
                print(xuehao,',',student_name,',', score,',',pingyu)
            '''

def parase_all_and_record_score(docx_path,student_idname,score_fiel):
    #打开学生名单excel，copy一份另存并写入成绩
    source_excel_file = xlrd.open_workbook(student_idname)
    dest_excel_file = copy(source_excel_file)
    read_table = source_excel_file.sheet_by_name(u'Sheet1')  # 通过名称获取
    write_table = dest_excel_file.get_sheet(0)

    max, min = get_max_min_size(docx_path)
    file_list = os.listdir(docx_path)

    #设置单元格格式，初始全部写入0和黄底
    pattern = xlwt.Pattern()  # Create the Pattern
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN  # May be: NO_PATTERN, SOLID_PATTERN, or 0x00 through 0x12
    pattern.pattern_fore_colour = 5  # May be: 8 through 63. 0 = Black, 1 = White, 2 = Red, 3 = Green, 4 = Blue, 5 = Yellow, 6 = Magenta, 7 = Cyan, 16 = Maroon, 17 = Dark Green, 18 = Dark Blue, 19 = Dark Yellow , almost brown), 20 = Dark Magenta, 21 = Teal, 22 = Light Gray, 23 = Dark Gray, the list goes on...
    style = xlwt.XFStyle()  # Create the Pattern
    style.pattern = pattern  # Add Pattern to Style

    #初始化所有人成绩为0
    for i in range(read_table.nrows):
        write_table.write(i, 3, 0, style)

    for i in range(read_table.nrows):
        write_table.write(i, 4, 0, style)

    student_namelist = []
    for file_name in file_list:
        file_name_list = []
        file_name_list = file_name.split('-')

        student_tmp = file_name_list[2].split(".")
        #student_name = student_tmp[0]  #名字所在字段不一样，取得位置也不一样

        student_name = file_name_list[1]


        student_namelist.append(student_name)
        score = 0
        find=0
        for i in range(read_table.nrows):
            get_name_fromcell = read_table.cell(i,2).value
            if(get_name_fromcell == student_name):
                #print(student_tmp)
                if student_tmp[1] == "docx":
                    score = get_score_from_file(docx_path+file_name, max, min)
                else:
                    print("Error:", student_tmp)
                #elif student_tmp == "doc":
                #    document = doc.Document(file_name)
                #    score = get_score_from_file(document, max, min)
                write_table.write(i, 3, score)

                write_table.write(i, 4, 1)
                find = 1
        if(find == 0):
            print("没有学生：",file_name, student_name)

    counter = collections.Counter(student_namelist)
    print(counter)
    dest_excel_file.save(score_fiel)

def parase_all_and_insert_db(docx_path):
    clear_table('words_temp')  # 清除数据表
    operatMySQl = OperateMySQL()

    jieba.load_userdict("dict.txt")

    # 获取文本内容
    contents_list = []
    words_list = []
    file_list = os.listdir(docx_path)
    for filename in file_list:
        filepathname= docx_path + filename
        print(filepathname)
        contents_list.clear()
        document = docx.Document(filepathname)
        #取文件内容
        for paragraph in document.paragraphs:
            contents_list.append(paragraph.text.encode('utf-8'))
        # 分词
        for text in contents_list:
            splits = jieba.cut(text)
            for word in splits:
                words_list.append(word.strip().lower())

    # 统计分词频率，计入数据库
    counter = collections.Counter(words_list)

    # 插入数据库
    for collect in counter:
        sqli = "insert into words_temp values ('{0}',{1});"
        sqlm = sqli.format(collect, counter[collect])
        operatMySQl.execute(sqlm)
    operatMySQl.commit()


    # 清洗无用数据
    clean_date_table = [ ' ', '\t', '\n', '、', '.', '，', '/', '；', '（', '）','。','：',',','-',';',')',':',
                        '的', '和','熟悉','开发','等','有','相关','及','能力','优先','工作','良好','与','对','年','常用',
                        '能够','了解','者','专业','熟练','能','计算机','专业','电子','进行',
                        '以上','具有','使用','精通','或','嵌入式','工程师',
                        ' ','1','2','3','4','5','6','7','8','9','0']
    sqli = "delete from words_temp  where words_content = '{0}';"
    for clean_item in clean_date_table:
        sqlm = sqli.format(clean_item)
        operatMySQl.execute(sqlm)
    operatMySQl.commit()


if __name__ == '__main__':
    starttime = datetime.datetime.now()
    print('Start time is %s.' % (str(datetime.datetime.now())))

    #插入数据库是5,6 和 7,8 只能各自处理，如果需要统计总的需要特殊处理
    #parase_all_and_insert_db(docx_path_56)
    #parase_all_and_record_score(docx_path_56,student_idname_56,score_fiel_56)

    #parase_all_and_insert_db(docx_path_78)
    #parase_all_and_record_score(docx_path_78,student_idname_78,score_fiel_78)

    #parase_all_and_insert_db(docx_path_all)

    #parase_all_and_insert_pingyu(docx_path_56,pingyu_56, score_pingyu_56)

    parase_all_and_insert_pingyu(docx_path_78, pingyu_56, score_pingyu_78)
    # 程序结束时间 及 耗时
    timedelta = datetime.datetime.now() - starttime
    print('End time is %s.' % (str(datetime.datetime.now())))
    print('Total test execution duration is %s.' % (timedelta.__str__()))



